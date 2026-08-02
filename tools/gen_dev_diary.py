"""生成《怪物森林》开发手记 docx。一次性脚本，跑完可删。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 全局中文字体
style = doc.styles["Normal"]
style.font.name = "PingFang SC"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x20, 0x40, 0x20)


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x30, 0x50, 0x30)


def p(text, bold=False, italic=False, size=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return para


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def code(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Menlo"
    run.font.size = Pt(9.5)
    para.paragraph_format.left_indent = Cm(0.5)


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = htext
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def hr():
    doc.add_paragraph("─" * 40)


# ---------- 标题 ----------
title = doc.add_heading("《怪物森林》开发手记：一个人 + 两层 AI，做完我的第一款「NPC 会活的」模拟经营", level=0)
p("主题：把「NPC 只会念固定台词」这件事，用 AI 真正解决掉。", italic=True)
p("工具链：Godot 引擎 · 舞台　CodeMaker · 编程搭档　DeepSeek · NPC 的大脑（运行时，不是开发期）", italic=True)

# ---------- 一、缘起 ----------
h1("一、缘起：为什么是「NPC 小镇」，为什么是现在")
p("我一直有个念头：想做一个不靠对话分支树、而是「你说过的话 TA 真记得」的小镇。"
  "以前这类念头只能停在设计文档里，卡在两件事上——策划写不完的对话分支；"
  "传统 NPC AI 没有语言能力，顶多是状态机加几句固定台词。现在不一样了：")
bullet("Godot 4.6 给了我一个人能 hold 住的舞台；")
bullet("CodeMaker 在前端 GDScript 和后端 Python 两套代码之间来回改，把系统一层层堆起来；")
bullet("DeepSeek LLM 把「一份人设 JSON」直接变成会看脸色说话、会记仇、会站队的居民，"
       "且是国产 API，跑一整个小镇的对话成本可控。")
p("于是有了《怪物森林》：一个九位居民都不是脚本 NPC 的模拟经营 + 竞选镇长的单机原型。"
  "玩家刷好感、送礼、打听八卦、拉票、参选镇长，镇上居民会为你拉票，也会在辩论台上跟你抢立场。")

# ---------- 二、项目一览 ----------
h1("二、项目一览")
table(
    ["维度", "数据"],
    [
        ["引擎", "Godot 4.6（GL Compatibility）"],
        ["客户端代码", "35 个 GDScript，约 8446 行（main.gd 1125 行、animal.gd 870 行、election_hud.gd 859 行、agent_client.gd 656 行）"],
        ["服务端代码", "32 个 Python 文件，约 11662 行（main.py 2351 行、agent.py 1730 行、election.py 1125 行、debate.py 609 行）"],
        ["数据配置", "18 份 JSON：9 位 NPC 人设、78 个任务、12 种危机、9 道辩论题（横跨 6 个议题）、地图/路网/刷新点等，全数据驱动"],
        ["美术与音乐", "复用免费 / CC0 素材包（Ninja Adventure、72 Character Free），未走 AI 生图 / 生乐这条线"],
        ["测试", "9 个测试文件，约 2140 行，回归靠 pytest 全量跑"],
        ["开发人力", "1 人（策划 + 全栈）"],
        ["开发节奏", "71 天（2026-05-23 首个 commit → 2026-08-02），92 次提交，1 个正式 tag（v0.3.0）"],
    ],
)

# ---------- 三、角色分工 ----------
h1("三、角色分工：一份代码，两层 AI")
p("跟很多「AI 辅助开发」项目不同，这里的 AI 分两层，一层管「怎么写出来」，一层管「游戏跑起来时怎么演」：")
bullet("开发期：CodeMaker 是编程搭档。一句「辩论结算面板跟增减记录对不上」，"
       "它会同时读 election.py / main.py / election_hud.gd 找根因再一起改，不是改一处留一堆引用错误。")
bullet("运行期：DeepSeek LLM 是「游戏里的演员」。每个 NPC 打招呼、送礼反应、造谣辟谣、辩论发言、"
       "对镇务任务的态度——全是运行时现场生成的台词，不是策划预先写好的文本库。")

code(
    "┌──────────────┐\n"
    "│ 我（策划/开发） │\n"
    "└──────┬───────┘\n"
    "       │ 需求 / Prompt / 取舍\n"
    "       ▼\n"
    "┌─────────────────────┐\n"
    "│      CodeMaker        │  开发期：改 .gd / .py，把系统一层层堆起来\n"
    "└──────────┬───────────┘\n"
    "           ▼\n"
    "┌─────────────────────────────┐\n"
    "│  Godot 4.6（前端舞台）        │\n"
    "│  场景树 · Autoload · 渲染/输入 │\n"
    "└──────────┬───────────────────┘\n"
    "           │ WebSocket\n"
    "           ▼\n"
    "┌─────────────────────────────┐\n"
    "│ Python 后端（世界状态机）      │\n"
    "│ FastAPI · 会话隔离 · 数值系统  │\n"
    "└──────────┬───────────────────┘\n"
    "           │ 每次开口现场调用\n"
    "           ▼\n"
    "┌─────────────────────────────┐\n"
    "│  DeepSeek LLM（运行期演员）    │\n"
    "│  9 位 NPC 的嘴，各有记忆与立场  │\n"
    "└─────────────────────────────┘\n"
    "           ▼\n"
    "       可玩的 Demo"
)

# ---------- 四、Godot ----------
h1("四、Godot：把「前端舞台」和「后端世界」缝在一起")

h2("4.1 场景树 + Autoload")
p("7 个 Autoload 单例：WorldClock（游戏内时钟）、LocationDB、AgentClient（WebSocket 客户端）、"
  "ChatManager、PlayerInventory、ItemDB、AudioManager。AgentClient 是前后端唯一通道——"
  "所有对话、送礼、投票、辩论、镇务指令都走它的一条 WebSocket。")

h2("4.2 GDScript 对 AI 足够友好")
p("跟很多用 AI 写 Godot 项目的经验一致：语法接近 Python、场景文件（.tscn）是纯文本 AI 也能读改、"
  "报错直接进 Output 面板，复制回去就能继续让 AI 修，改代码这件事能一直循环下去而不会散架。")

h2("4.3 动态 NPC 实例化：从写死节点到按名单建人")
p("最初 main.tscn 里 6 个 NPC 是写死的节点。后来做「任期结束换一个村民」的轮换玩法时发现："
  "写死的节点没法反映「熟人搬走、生人搬进来」这件事，改成全动态：后端下发在场名单（roster），"
  "前端按名单实例化/删除节点、改门牌。这一步顺带牵出一个很隐蔽的架构坑（见第八节）。")

# ---------- 五、Python 后端 ----------
h1("五、Python 后端：把「小镇」写成一台真正的状态机")
p("这是跟纯前端 + AI 编程助手项目最大的结构差异——NPC 要记事、要投票、要有立场，"
  "这类状态没法只放在前端本地存档里，得有一个持续运行、能断线重连的世界。")

h2("5.1 会话隔离差点被我自己破坏")
p("架构前提是「每个玩家（sid）一份独立 SQLite，personas / LLM 全局共享」。"
  "轮换新增 3 个备选 NPC 时，第一版做法是「按当前会话的在场名单裁剪全局 personas」——"
  "结果是第一个连上来的玩家的轮换结果，污染了之后所有新玩家的开局。"
  "退回「personas 全量常驻，运行时按会话的 present_provider 回调过滤」才是这套多会话架构的本意。")

h2("5.2 十余个业务模块，边界互不越权")
bullet("election：选举权重、唱票演出、竞选分增减记录")
bullet("debate：阵营辩论、议题站位、信息差")
bullet("crisis：镇务危机事件")
bullet("rumor：造谣 / 打听 / 辟谣")
bullet("power / mayor_tasks：镇长权力点行动、镇务派工")
bullet("roster：NPC 任期轮换（走 1 进 1，恒定在场 6 人）")
bullet("affection / items / gifts：好感度、送礼经济、道具可达性")
p("硬规矩是「谁的数值谁说了算」——比如好感该涨该跌，只有 affection.py 一处能改，"
  "选举/辩论/镇务模块都只能调它，不能自己动好感值。")

h2("5.3 数据驱动")
p("人设、任务、危机、辩论题全放 JSON，改数值不用碰代码——策划在表里改一列，游戏里就变了。")

# ---------- 六、DeepSeek LLM ----------
h1("六、DeepSeek LLM：让 NPC 有嘴，但不让它管数字")
p("这是从多次真实 bug 里换来的设计铁律：LLM 只负责「怎么说」，绝不让它决定「该涨多少好感 / 该不该给东西」。"
  "所有数值结果——好感增减、选票权重、任务成功率、送礼库存——全部由确定性公式算好，"
  "LLM 只在算好的结果上现场编一句台词，且每处都留兜底模板：LLM 超时或挂了，游戏也不能卡在「正在思考」。")
p("举例：greet() 先由后端算出「是否认得玩家」「玩家是不是现任镇长」，再把这个前提喂给 LLM，"
  "不是让它自由发挥「要不要认识你」；镇长索取道具时说的那句不情愿的嘀咕、辩论台词、拉票效果——"
  "判定逻辑全在 Python 里，LLM 只出嘴，不判生死。")
p("技术上走 OpenAI 兼容 SDK 调 DeepSeek（deepseek-v4-flash），国产 API 价格能撑住一整个小镇的高频对话。")

# ---------- 七、美术与音乐 ----------
h1("七、美术与音乐：这次没让 AI 画，是故意的")
p("跟很多同类分享不同，这个项目的美术资源全部来自开源 / 免费素材包（Ninja Adventure、72 Character Free 等，"
  "CC0 或类似协议），没有走 AI 生图 / 生乐这条线。原因很直接：这个项目的含金量在「NPC 大脑」，不在画面——"
  "把省下来的时间全砸进玩法迭代（选举、辩论、拉票、镇务）比多磨一版像素画更划算。")
p("9 位 NPC 的立绘/头像、地图 tileset、188 条 BGM/SFX，全部拿来即用，靠 Godot 的 .import 管线统一接入，"
  "零额外制作成本。")

# ---------- 八、踩过的坑 ----------
h1("八、踩过的坑（这次是真事故，不是软文）")

p("1. 疲劳度符号反转刷分", bold=True)
p("送礼疲劳系数本意是「同一件东西送多了不新鲜」，但负分礼物（讨厌的东西）乘上负疲劳系数会变成正——"
  "「同一句难听话骂多了反而变夸奖」。修法：疲劳系数下钳到 0，只能削弱惩罚，不能反转符号。")

p("2. 回礼池「替换」而非「累加」", bold=True)
p("某 NPC 的低档回礼只配在 common 档，好感升到 love 后池子被高档整个替换掉——"
  "关系越好反而越拿不到低档货，完全反直觉。改成按档位加权累加抽取，低档永远保留。")

p("3. 多会话架构差点被自己人破坏", bold=True)
p("见第五节：裁剪全局 personas 差点让所有玩家共享同一个人的轮换结果。")

p("4. NPC 把自己讲成八卦主角", bold=True)
p("小蓝对玩家说「那位旅人给小蓝送了挺贵重的东西」——排除逻辑只挡了「谁在说」，"
  "没挡「内容里提到了谁」。加一列 involved_ids 记录内容里被卷进来的 NPC 才补全。")

p("5. 送出自己根本没有的东西", bold=True)
p("焰仔一边发着「找一件火卷轴」的任务，一边说「我这有」就送了出去——"
  "loves（喜欢收到什么）被误当成 giveable（手上有什么）。"
  "更糟的是这打通了一条「向他要走再送回去」白刷好感的闭环。")

# ---------- 九、结语 ----------
h1("九、结语：AI 没有替我做游戏，是 CodeMaker 和 DeepSeek 一起，让「NPC 是活人」这件事真正跑起来")
p("过去这类「NPC 有记忆、有立场」的念头只能停在设计文档——受限于人力（写不完的对话分支树）"
  "和技术（传统 NPC AI 没有语言能力）。这次 CodeMaker 接管了「把系统一层层堆起来」这件苦力活，"
  "DeepSeek 接管了「每个 NPC 该怎么说话」这件以前只能靠海量分支树硬堆的活，"
  "我则从「策划 + 纯手写代码」变成「导演 + 判断官 + 数值配平师」。")

p("下一步想做的事：", bold=True)
bullet("备选池里另外两位新村民（夜行作息的老土、跟他关系 -30 的莓莓）还没真正登场验证过"
       "「作息摩擦」和「调解难度」这两个设计意图；")
bullet("数值手感（衰减曲线、难度爬升、挖墙脚幅度）目前只验证了「算对了」，没验证"
       "「玩起来是不是那个手感」，得连打几届才知道；")
bullet("想把「镇长」这个身份的权力感再做实一点——这轮刚加了「索取」和「命令」两个新动作，"
       "长线还想让镇务玩法承担更多「后期无所事事」问题的解药。")

hr()
p("感谢 Godot、CodeMaker、DeepSeek，也感谢这个能让一个人把一整个镇的活儿干完的时代。", italic=True)

doc.save("/Users/lddx/Documents/GitHub/game01/build/怪物森林开发手记.docx")
print("DONE, total paragraphs:", len(doc.paragraphs))
